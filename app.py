# -*- coding: utf-8 -*-
# 🚀 晋江市第八实验小学AI教案系统 v10.3 (豪华完整版)
import streamlit as st
from datetime import datetime
import asyncio
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, List, Optional, Tuple
import logging
import io
import sys
from sparkai.llm.llm import ChatSparkLLM, ChunkPrintHandler
from sparkai.core.messages import ChatMessage

# ==================== 🔧 配置初始化 ====================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('ai_lesson_plan.log'),
        logging.StreamHandler()
    ]
)

# ==================== 🏫 教育配置模块 ====================
class EduConfig:
    """教学资源配置中心"""
    SCHOOL_NAME = "晋江市第八实验小学"
    SCHOOL_MOTTO = "让教育成为孩子生命中的美好印记"
    LOGO_URL = "https://img.icons8.com/clouds/500/school.png"
    
    SUBJECTS = ["语文", "数学", "英语", "科学", "道德与法治", "信息技术"]
    GRADES = ["一年级", "二年级", "三年级", "四年级", "五年级", "六年级"]
    
    @staticmethod
    def generate_5e_template(topic: str) -> str:
        """5E教学法模板"""
        return f"""
        # 《{topic}》5E教学设计
        ## 参与(Engagement)
        - 情境创设：通过______激发学生兴趣
        - 问题提出：______
        
        ## 探索(Exploration)
        - 小组活动：______
        - 实验设计：______
        
        ## 解释(Explanation)
        - 核心概念：______
        - 师生对话：______
        
        ## 深化(Elaboration)
        - 跨学科联系：______
        - 现实应用：______
        
        ## 评价(Evaluation)
        - 形成性评价：______
        - 总结性评价：______
        """

    TEACHING_METHODS = {
        "5E教学法": generate_5e_template,
        "项目式学习(PBL)": lambda t: f"# 《{t}》项目式学习\n1. 驱动问题：______\n2. 成果展示：______",
        "翻转课堂": lambda t: f"# 《{t}》翻转课堂\n1. 课前微课：______\n2. 课堂活动：______",
        "情境教学法": lambda t: f"# 《{t}》情境教学\n1. 情境创设：______\n2. 角色扮演：______",
        "探究式学习": lambda t: f"# 《{t}》探究式学习\n1. 问题提出：______\n2. 实验探究：______",
        "自由探索": lambda t: f"# 《{t}》自由探索\n"
    }

# ==================== 🤖 AI核心模块 ==================== 
class SparkAI:
    """工业级AI引擎"""
    MAX_RETRIES = 3

    @staticmethod
    async def _call_api(prompt: str) -> Tuple[Dict, str]:
        """执行API调用"""
        last_error = ""
        for attempt in range(1, SparkAI.MAX_RETRIES + 1):
            try:
                SPARKAI_URL = 'wss://spark-api.xf-yun.com/v4.0/chat'
                SPARKAI_APP_ID = 'df9b6f91'
                SPARKAI_API_SECRET = 'ZDY4NDQ5NGM2ODg1MjRiYTliYzQxODU5'
                SPARKAI_API_KEY = '7965665a57e9554153c6570694ab50e7'
                SPARKAI_DOMAIN = '4.0Ultra'
                
                spark = ChatSparkLLM(
                    spark_api_url=SPARKAI_URL,
                    spark_app_id=SPARKAI_APP_ID,
                    spark_api_key=SPARKAI_API_KEY,
                    spark_api_secret=SPARKAI_API_SECRET,
                    spark_llm_domain=SPARKAI_DOMAIN,
                    streaming=False,
                )
                messages = [ChatMessage(role="user", content=prompt)]
                handler = ChunkPrintHandler()
                response = spark.generate([messages], callbacks=[handler])
                
                return {
                    "payload": {
                        "choices": [{
                            "text": [{
                                "content": response.generations[0][0].text
                            }]
                        }]
                    }
                }, ""
            except Exception as e:
                last_error = f"尝试 {attempt}/{SparkAI.MAX_RETRIES}: {str(e)}"
                if attempt < SparkAI.MAX_RETRIES:
                    await asyncio.sleep(1)
        return {}, f"所有尝试失败: {last_error}"

    @staticmethod
    async def generate_lesson(topic: str, method: str) -> Tuple[str, Optional[str]]:
        """生成教案内容"""
        template = EduConfig.TEACHING_METHODS.get(method, EduConfig.generate_5e_template)(topic)
        prompt = f"作为资深教师，请完善以下{method}教案（用中文回答）：\n{template}"
        response, error = await SparkAI._call_api(prompt)
        if error:
            logging.warning(f"API调用失败: {error}")
            return template, error
        try:
            content = response.get("payload", {}).get("choices", [{}])[0].get("text", [{}])[0].get("content", "")
            return content.strip() if content.strip() else template, None
        except Exception as e:
            logging.error(f"响应解析错误: {str(e)}")
            return template, str(e)

# ==================== 📄 完整文档生成模块 ====================
# -*- coding: utf-8 -*-
# 🏫 晋江市第八实验小学文档生成模块 (终极完整版)
import io
import logging
from datetime import datetime
from typing import Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocBuilder:
    """专业文档生成器 (完整保障版)"""

    @staticmethod
    def create_doc(content: str, meta: Dict) -> io.BytesIO:
        """
        生成完整教案文档
        Args:
            content: Markdown格式的教案内容
            meta: 包含学科/年级等信息的字典
                必需字段: topic, subject, grade, method
        Returns:
            可直接下载的文档内存流
        """
        # === 文档构建流程 ===
        doc = Document()
        
        # 1. 设置全局样式
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        # 2. 添加学校封面页
        header = doc.sections[0].header
        header.paragraphs[0].text = "晋江市第八实验小学官方教案"
        
        title = doc.add_heading(level=0)
        title_run = title.add_run("晋江市第八实验小学\n")
        title_run.font.size = Pt(26)
        title_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 3. 添加教案基本信息
        doc.add_heading(f"《{meta.get('topic', '')}》教学方案", level=1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Shading Accent 1'
        info_table.cell(0, 0).text = "学科/年级"
        info_table.cell(0, 1).text = f"{meta.get('subject', '')} {meta.get('grade', '')}"
        info_table.cell(1, 0).text = "教学方法"
        info_table.cell(1, 1).text = meta.get('method', '')
        info_table.cell(2, 0).text = "设计教师"
        info_table.cell(2, 1).text = "________________"
        info_table.cell(3, 0).text = "生成时间" 
        info_table.cell(3, 1).text = datetime.now().strftime('%Y年%m月%d日')

        # 4. 添加教案内容
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            elif line.startswith('1. '):
                p = doc.add_paragraph(style='List Number')
                p.add_run(line[3:])
            else:
                doc.add_paragraph(line)

        # 5. 添加页脚
        footer = doc.sections[0].footer
        footer.paragraphs[0].text = f"晋江市第八实验小学 ∙ {datetime.now().year}"
        footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 生成文件流
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

# ==================== 🏫 无Logo但保留学校信息的豪华版 ====================
class UIManager:
    """移除Logo但保留学校其他信息的专业界面"""
    
    @staticmethod
    def _inject_styles():
        """注入专业级CSS样式"""
        st.markdown(f"""
        <style>
            /* 主色调 */
            :root {{
                --primary: #4361ee;
                --secondary: #3a0ca3;
                --accent: #f72585;
            }}
            
            /* 学校标题区 */
            .school-header {{
                background: linear-gradient(135deg, var(--primary), var(--secondary));
                color: white;
                padding: 2.5rem;
                border-radius: 12px;
                margin-bottom: 2rem;
                text-align: center;
                box-shadow: 0 4px 20px rgba(67, 97, 238, 0.15);
                position: relative;
                overflow: hidden;
            }}
            
            /* 动态装饰元素 */
            .school-header::before {{
                content: "";
                position: absolute;
                top: -50%;
                left: -50%;
                width: 200%;
                height: 200%;
                background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, rgba(255,255,255,0) 70%);
                animation: pulse 8s infinite linear;
            }}
            
            @keyframes pulse {{
                0% {{ transform: rotate(0deg); }}
                100% {{ transform: rotate(360deg); }}
            }}
            
            /* 专业卡片 */
            .pro-card {{
                background: white;
                border-radius: 12px;
                padding: 1.5rem;
                box-shadow: 0 4px 12px rgba(0,0,0,0.08);
                margin-bottom: 1.5rem;
                border-left: 4px solid var(--accent);
                transition: all 0.3s;
            }}
        </style>
        """, unsafe_allow_html=True)

    @staticmethod
    def setup():
        """初始化界面（保留学校文字信息）"""
        st.set_page_config(

page_title=f"{EduConfig.SCHOOL_NAME}AI教案系统",
            page_icon="📚",
            layout="wide"
        )
        UIManager._inject_styles()
        
        # 豪华标题区（用文字校徽替代图片）
        st.markdown(f"""
        <div class="school-header">
            <h1 style="margin-bottom: 0.5rem; font-size: 2.2rem;">
                {EduConfig.SCHOOL_NAME}AI教案系统
            </h1>
            <p style="font-size: 1.1rem; opacity: 0.9; margin-bottom: 0;">
                {EduConfig.SCHOOL_MOTTO}
            </p >
            <div style="margin-top: 1rem;">
                <span style="
                    display: inline-block;
                    background: rgba(255,255,255,0.2);
                    padding: 0.25rem 1rem;
                    border-radius: 20px;
                    font-size: 0.9rem;
                ">
                🎯 专业版 v10.3
                </span>
            </div>
        </div>
        """, unsafe_allow_html=True)

    @staticmethod
    def show_controls() -> Optional[Dict]:
        """显示控制面板（专业无Logo版）"""
        with st.sidebar:
            # 配置面板标题
            st.markdown(f"""
            <div style="
                text-align: center;
                margin-bottom: 1.5rem;
                padding-bottom: 0.5rem;
                border-bottom: 1px solid #eee;
            ">
                <h3>⚙️ {EduConfig.SCHOOL_NAME}教案配置</h3>
            </div>
            """, unsafe_allow_html=True)
            
            with st.form("lesson_form"):
                # 表单元素
                topic = st.text_input(
                    "📝 教学主题*",
                    placeholder=f"例：{EduConfig.SCHOOL_NAME}特色课程",
                    help="请输入本节课的教学主题"
                )
                
                col1, col2 = st.columns(2)
                with col1:
                    subject = st.selectbox(
                        "📚 学科*", 
                        EduConfig.SUBJECTS,
                        index=0
                    )
                with col2:
                    grade = st.selectbox(
                        "🎒 年级*",
                        EduConfig.GRADES,
                        index=2
                    )
                
                method = st.selectbox(
                    "🏫 教学方法",

list(EduConfig.TEACHING_METHODS.keys()),
                    index=0
                )
                
                submitted = st.form_submit_button(
                    "✨ 生成专业教案", 
                    type="primary",
                    use_container_width=True
                )
                
                if submitted:
                    if not topic.strip():
                        st.error("请输入教学主题！", icon="🚨")
                        return None
                    return {
                        "topic": topic,
                        "subject": subject,
                        "grade": grade,
                        "method": method
                    }
        return None

    @staticmethod
    def show_result(content: str, error: Optional[str], config: Dict):
        """显示结果（豪华无Logo版）"""
        # 错误处理
        if error:
            st.error(f"⛔ 生成错误: {error}", icon="⚠️")
            st.warning("已使用基础模板生成内容", icon="ℹ️")
    
        # 标签页布局
        tab1, tab2 = st.tabs(["📖 教案预览", "📥 文档导出"])
    
        with tab1:
            st.markdown(f"""
            <div class="pro-card">
                <div style="
                    color: #4361ee;
                    font-weight: bold;
                    font-size: 1.1rem;
                    margin-bottom: 1rem;
                ">
                    {EduConfig.SCHOOL_NAME} {config['grade']}{config['subject']}教案
                </div>
                <div style="color: #666; margin-bottom: 0.5rem;">
                    主题：{config['topic']} | 方法：{config['method']}
                </div>
                {content}
            </div>
            """, unsafe_allow_html=True)
    
        with tab2:
            st.markdown(f"""
            <div class="pro-card">
                <h4 style="
                    color: #4361ee;
                    margin-top: 0;
                    border-left: 3px solid #f72585;
                    padding-left: 0.8rem;
                ">
                    ⬇️ {EduConfig.SCHOOL_NAME}教案文档导出
                </h4>
            </div>
            """, unsafe_allow_html=True)
        
            # 自动生成并触发下载
            try:
                doc_bytes = DocBuilder.create_doc(content, config)
                download_key = f"download_{hash(str(config))}"
            
                st.download_button(
                    label="📥 立即下载",
                    data=doc_bytes,
                    file_name=f"{EduConfig.SCHOOL_NAME}_{config['topic']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=download_key
                )
            
                # 自动触发下载
                st.markdown(f"""
                <script>
                    function autoClick() {{
                        const button = parent.document.querySelector('button[data-testid="baseButton-secondary"][aria-label="{download_key}"]');
                        if (button) {{
                            button.click();
                        }} else {{
                            setTimeout(autoClick, 200);
                        }}
                    }}
                    setTimeout(autoClick, 300);
                </script>
                """, unsafe_allow_html=True)
            
            except Exception as e:
                st.error(f"文档生成失败: {str(e)}")
        
            st.caption(f"💡 提示：下载文档将包含{EduConfig.SCHOOL_NAME}官方页眉页脚")


# ==================== 🚀 豪华主程序 ====================
async def main_async():
    """豪华版主程序"""
    try:
        # 初始化豪华界面
        UIManager.setup()
        
        # 显示控制面板
        config = UIManager.show_controls()
        
        if config:
            # 豪华生成过程
            with st.spinner("🛠️ 正在准备教学资源..."):
                await asyncio.sleep(0.5)
            
            progress_bar = st.progress(0)
            status_container = st.empty()
            
            phases = [
                ("📚 分析教学大纲", 20),
                ("🧩 设计课堂活动", 40),
                ("✍️ 编写教案内容", 60),
                ("🔍 质量检查", 80),
                ("🎉 完成生成", 100)
            ]
            
            content, error = "", None
            for phase_name, progress in phases:
                status_container.markdown(
                    f"""<div style="background:#f0f2f6;padding:15px;border-radius:10px">
                    <h4 style="color:#2e75b6">⏳ {phase_name}...</h4>
                    <p style="font-size:12px;color:#666">正在生成 {config['grade']}{config['subject']}《{config['topic']}》教案</p >
                    </div>""",
                    unsafe_allow_html=True
                )
                progress_bar.progress(progress)
                await asyncio.sleep(0.8)
                
                if progress == 60:
                    content, error = await SparkAI.generate_lesson(
                        topic=config["topic"],
                        method=config["method"]
                    )
            
            progress_bar.empty()
            status_container.empty()
            
            # 显示结果
            UIManager.show_result(content, error, config)

    except Exception as e:
        st.error(f"系统发生错误: {str(e)}", icon="🚨")
        logging.error(f"系统错误: {str(e)}", exc_info=True)

def main():
    # ============== 防护代码开始 ==============
    import streamlit as st
    from datetime import datetime
    
    # 1. 时间限制（8:00-18:00开放）
    if not datetime.now().hour in range(8, 18):
        st.error("⏰ 系统开放时间：每日8:00-18:00", icon="🚫")
        st.stop()
    
    # 2. 密码验证
    if not hasattr(st.session_state, 'auth'):
        with st.form("auth_form"):
            st.markdown("### " + EduConfig.SCHOOL_NAME + "访问验证")
            pwd = st.text_input("请输入管理员密码", type="password")
            if st.form_submit_button("授权访问"):
                if pwd == "8888":  # 请修改为您的密码
                    st.session_state.auth = True
                    st.rerun()
        st.stop()
    # ============== 防护代码结束 ==============

    # 原有代码保持不变
    if sys.platform == "win32":
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        loop.run_until_complete(main_async())
    finally:
        loop.close()

if __name__ == "__main__":
    main()